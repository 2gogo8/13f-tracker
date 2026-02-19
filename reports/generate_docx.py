#!/usr/bin/env python3
"""
JG反市場報告書 - Markdown to DOCX Converter
Converts the markdown report to a professional Word document with custom formatting.
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_number(section):
    """Add page numbers to the footer"""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page number field
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    run.font.size = Pt(10)
    run.font.name = 'Arial'


def setup_document_styles(doc):
    """Configure document styles with Helvetica/Arial fonts"""
    # Set up default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(20)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 3 style
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0, 0, 0)
    
    return doc


def setup_page_margins(doc):
    """Set 1-inch margins on all sides"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_height = Inches(11.69)  # A4 height
        section.page_width = Inches(8.27)    # A4 width


def create_title_page(doc):
    """Create the title page with specified content"""
    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('JG的反市場報告書')
    run.font.name = 'Arial'
    run.font.size = Pt(28)
    run.font.bold = True
    
    # Add spacing
    doc.add_paragraph()
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('深度研究系列 — 五檔反市場精選個股')
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    
    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run('2026年2月19日')
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    
    # Page break after title page
    doc.add_page_break()


def create_toc_page(doc):
    """Create table of contents page"""
    # TOC title
    toc_title = doc.add_heading('目錄 (Table of Contents)', level=1)
    
    # TOC entries
    companies = [
        ('GOOGL', '第 3 頁'),
        ('PLTR', '第 19 頁'),
        ('RKLB', '第 30 頁'),
        ('AMD', '第 42 頁'),
        ('ARES', '第 53 頁')
    ]
    
    # Create table
    table = doc.add_table(rows=len(companies) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '公司'
    hdr_cells[1].text = '頁碼'
    
    # Data rows
    for i, (company, page) in enumerate(companies, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = company
        row_cells[1].text = page
    
    # Style table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
    
    # Add note
    doc.add_paragraph()
    note = doc.add_paragraph('本版新增：資本支出應用場景白話解釋、每家公司競爭對手比較（2026差異）、投資結論改為『最大優勢/最大難處』。')
    note_run = note.runs[0]
    note_run.font.name = 'Arial'
    note_run.font.size = Pt(10)
    note_run.font.italic = True
    
    # Page break after TOC
    doc.add_page_break()


def parse_inline_formatting(text):
    """Parse markdown bold (**text**) and italic (*text*) into runs"""
    runs = []
    pattern = r'(\*\*.*?\*\*|\*.*?\*|[^*]+|\*)'
    
    for match in re.finditer(pattern, text):
        segment = match.group()
        
        if segment.startswith('**') and segment.endswith('**'):
            # Bold text
            runs.append(('bold', segment[2:-2]))
        elif segment.startswith('*') and segment.endswith('*') and not segment.startswith('**'):
            # Italic text
            runs.append(('italic', segment[1:-1]))
        elif segment:
            # Normal text
            runs.append(('normal', segment))
    
    return runs


def add_formatted_paragraph(doc, text, style='Normal'):
    """Add paragraph with inline formatting (bold/italic)"""
    para = doc.add_paragraph(style=style)
    
    # Parse and add runs with formatting
    formatted_runs = parse_inline_formatting(text)
    
    for format_type, content in formatted_runs:
        run = para.add_run(content)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        
        if format_type == 'bold':
            run.font.bold = True
        elif format_type == 'italic':
            run.font.italic = True
    
    return para


def parse_markdown_table(lines, start_idx):
    """Parse markdown table and return table data and ending index"""
    table_lines = []
    idx = start_idx
    
    # Collect table lines
    while idx < len(lines) and '|' in lines[idx]:
        table_lines.append(lines[idx])
        idx += 1
    
    if len(table_lines) < 2:
        return None, start_idx
    
    # Parse header and rows
    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
    
    # Skip separator line
    data_lines = table_lines[2:]
    
    rows = []
    for line in data_lines:
        if not line.strip():
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    
    return (headers, rows), idx


def add_markdown_table(doc, headers, rows):
    """Add a table to the document"""
    if not headers or not rows:
        return
    
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.bold = True
    
    # Add data rows
    for i, row_data in enumerate(rows, 1):
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].text = cell_text
                for paragraph in row_cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)


def process_markdown_content(doc, md_file_path):
    """Process markdown content and convert to Word document"""
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    skip_first_title = True  # Skip the first title as we have custom title page
    company_sections = ['GOOGL', 'PLTR', 'RKLB', 'AMD', 'ARES']
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip page markers
        if line.startswith('--- Page'):
            i += 1
            continue
        
        # Skip the first title (we have custom title page)
        if skip_first_title and line.startswith('#'):
            skip_first_title = False
            i += 1
            continue
        
        # Skip separator lines
        if line.strip() == '================================================================':
            i += 1
            continue
        
        # Skip empty lines at start
        if not line.strip():
            i += 1
            continue
        
        # Check for company section headers (add page break before)
        for company in company_sections:
            if f'第一部：{company}' in line or f'第二部：{company}' in line or \
               f'第三部：{company}' in line or f'第四部：{company}' in line or \
               f'第五部：{company}' in line:
                doc.add_page_break()
                break
        
        # Headers
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        
        # Bullet lists
        elif line.startswith('- ') or line.startswith('▸ '):
            # Remove bullet marker
            text = line[2:].strip()
            para = add_formatted_paragraph(doc, text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            para = add_formatted_paragraph(doc, text, style='List Number')
        
        # Tables
        elif line.strip().startswith('|'):
            table_data, next_idx = parse_markdown_table(lines, i)
            if table_data:
                headers, rows = table_data
                add_markdown_table(doc, headers, rows)
                i = next_idx - 1
        
        # Horizontal rules
        elif line.strip().startswith('---') and not 'Page' in line:
            doc.add_paragraph('_' * 50)
        
        # Regular paragraphs
        elif line.strip():
            add_formatted_paragraph(doc, line)
        
        i += 1


def main():
    """Main function to generate the DOCX file"""
    print("Starting document generation...")
    
    # Create document
    doc = Document()
    
    # Setup
    setup_document_styles(doc)
    setup_page_margins(doc)
    
    # Add page numbers to all sections
    for section in doc.sections:
        add_page_number(section)
    
    # Create title page
    print("Creating title page...")
    create_title_page(doc)
    
    # Create TOC page
    print("Creating table of contents...")
    create_toc_page(doc)
    
    # Process markdown content
    print("Processing markdown content...")
    md_file = '/Users/jgtruestock/.openclaw/workspace/projects/13f-tracker/reports/JG反市場報告書_完整文字版_頁碼版.md'
    process_markdown_content(doc, md_file)
    
    # Save document
    output_file = '/Users/jgtruestock/.openclaw/workspace/projects/13f-tracker/reports/JG反市場報告書_完整文字版.docx'
    print(f"Saving document to {output_file}...")
    doc.save(output_file)
    
    print("✅ Document generated successfully!")
    
    # Report file size
    import os
    file_size = os.path.getsize(output_file)
    print(f"📄 File size: {file_size:,} bytes ({file_size / 1024 / 1024:.2f} MB)")


if __name__ == '__main__':
    main()
